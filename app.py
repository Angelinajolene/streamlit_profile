import streamlit as st
from PIL import Image
import os
import PyPDF2
from docx import Document
import pandas as pd
import base64
import sqlite3
import re
import time

st.set_page_config(
    page_title="AI HR Employee Management System",
    page_icon="💼",
    layout="wide"
)

st.markdown("""
<style>
.main {
    background-color: #f5f7fa;
}
.title {
    text-align: center;
    color: #1f3c88;
    font-size: 40px;
    font-weight: bold;
    margin-bottom: 20px;
}
.box {
    background-color: white;
    padding: 25px;
    border-radius: 15px;
    box-shadow: 0px 2px 12px rgba(0,0,0,0.1);
    margin-bottom: 20px;
}
.stButton > button {
    width: 100%;
    border-radius: 10px;
    height: 45px;
    font-size: 16px;
    font-weight: bold;
}
.viewer-box {
    background-color: white;
    padding: 25px;
    border-radius: 15px;
    box-shadow: 0px 2px 12px rgba(0,0,0,0.1);
}
</style>
""", unsafe_allow_html=True)

# Database Connection
conn = sqlite3.connect("employees.db", check_same_thread=False)
cursor = conn.cursor()

cursor.execute('''
CREATE TABLE IF NOT EXISTS employees(
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT,
    age INTEGER,
    department TEXT,
    email TEXT,
    skills TEXT,
    image TEXT,
    resume TEXT
)
''')
conn.commit()

# =================================================
# SECURE ADMINISTRATIVE CREDENTIALS
# =================================================
ADMIN_EMAIL = "admin@company.com"
ADMIN_PASS = "admin1234"  

HR_EMAIL = "hr@company.com"
HR_PASS = "hr1234"        

# Session State Initialization
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user_role" not in st.session_state:
    st.session_state.user_role = None  
if "employee_id" not in st.session_state:
    st.session_state.employee_id = None  
if "selected_profile" not in st.session_state:
    st.session_state.selected_profile = None
if "selected_resume" not in st.session_state:
    st.session_state.selected_resume = None
if "open_edit_page" not in st.session_state:
    st.session_state.open_edit_page = False
if "edit_employee_id" not in st.session_state:
    st.session_state.edit_employee_id = None

def login():
    st.markdown("<h1 class='title'> Employee Login Portal </h1>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        st.markdown("<div class='box'>", unsafe_allow_html=True)
        
        username = st.text_input("Username ")
        email_input = st.text_input(" Email Address")
        
        # Smart UI Logic: Checks the email input dynamically as the user types
        is_management = email_input.strip().lower() in [ADMIN_EMAIL.lower(), HR_EMAIL.lower()]
        
        if is_management:
            # This password block pops up dynamically ONLY for Admin or HR emails
            st.warning("🔒 Management Email Detected. Password Required.")
            password_input = st.text_input("Enter Management Password", type="password")
        else:
            password_input = ""
        
        if st.button("Log In"):
            if not username.strip() or not email_input.strip():
                st.error("Please fill in both Username and Email fields to log in.")
            
            # 1. Secure Admin Check
            elif email_input.strip().lower() == ADMIN_EMAIL.lower():
                if password_input == ADMIN_PASS:
                    st.session_state.logged_in = True
                    st.session_state.user_role = "admin"
                    st.success(f"Admin Login Successful. Welcome back, {username}!")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("Access Denied: Incorrect password provided for Admin access.")
                
            # 2. Secure HR Check
            elif email_input.strip().lower() == HR_EMAIL.lower():
                if password_input == HR_PASS:
                    st.session_state.logged_in = True
                    st.session_state.user_role = "hr"
                    st.success(f"HR Login Successful. Welcome back, {username}!")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("Access Denied: Incorrect password provided for HR access.")
                
            # 3. Standard Worker Database Check
            else:
                cursor.execute("SELECT id, name FROM employees WHERE email=?", (email_input.strip(),))
                user_match = cursor.fetchone()
                
                if user_match:
                    st.session_state.logged_in = True
                    st.session_state.user_role = "employee"
                    st.session_state.employee_id = user_match[0]
                    st.success(f"Employee Login Successful! Welcome back, {user_match[1]}.")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("Access Denied. No matching Email address found in the database records.")
                    
        st.markdown("</div>", unsafe_allow_html=True)

def logout():
    st.session_state.logged_in = False
    st.session_state.user_role = None
    st.session_state.employee_id = None
    st.session_state.open_edit_page = False
    st.session_state.edit_employee_id = None
    st.session_state.selected_profile = None
    st.session_state.selected_resume = None
    st.rerun()

def extract_pdf_text(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted
    return text

def extract_docx_text(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def extract_skills(text):
    skill_keywords = [
        "Python", "Java", "SQL", "Machine Learning", "Data Science", 
        "HTML", "CSS", "JavaScript", "React", "Streamlit", "C", "C++", 
        "Power BI", "Excel"
    ]
    found_skills = []
    for skill in skill_keywords:
        if skill.lower() in text.lower():
            found_skills.append(skill)
    return found_skills

def show_pdf(file_path):
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode("utf-8")
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="700" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

def fetch_employees():
    cursor.execute("SELECT * FROM employees")
    return cursor.fetchall()

def fetch_single_employee(emp_id):
    cursor.execute("SELECT * FROM employees WHERE id=?", (emp_id,))
    return cursor.fetchone()

# =================================================
# EDIT PROFILE VIEW
# =================================================
def render_edit_form(employee_id, return_to_role_view):
    st.markdown("<h1 class='title'> Edit Profile Details</h1>", unsafe_allow_html=True)
    selected_employee = fetch_single_employee(employee_id)
    
    if selected_employee:
        full_name = selected_employee[1]
        current_age = selected_employee[2]
        current_department = selected_employee[3]
        current_email = selected_employee[4]
        current_skills = selected_employee[5]
        current_image = selected_employee[6]
        current_resume = selected_employee[7]
        
        name_parts = full_name.split()
        first_name_value = name_parts[0] if len(name_parts) > 0 else ""
        last_name_value = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""
        
        col1, col2 = st.columns([1,2])
        with col1:
            st.subheader("Current Profile Image")
            if current_image and os.path.exists(current_image):
                st.image(current_image, width=250)
            else:
                st.info("No profile picture uploaded.")
                
            uploaded_new_image = st.file_uploader("Upload New Profile Picture", type=["jpg", "jpeg"])
            
        with col2:
            new_first_name = st.text_input("First Name", value=first_name_value)
            new_last_name = st.text_input("Last Name", value=last_name_value)
            new_age = st.number_input("Age", min_value=18, max_value=60, value=int(current_age))
            
            department_options = ["IT", "HR", "Finance", "Marketing"]
            if current_department not in department_options:
                department_options.append(current_department)
                
            new_department = st.selectbox("Department", department_options, index=department_options.index(current_department))
            new_email = st.text_input("Email", value=current_email)
            
            st.subheader("Skills Asset Track")
            st.info(current_skills if current_skills else "No skills parsed yet.")
            
            uploaded_new_resume = st.file_uploader("Upload New Resume", type=["pdf", "docx"])
            
            updated_resume_path = current_resume
            updated_image_path = current_image
            updated_skills = current_skills
            
            if uploaded_new_resume is not None:
                if uploaded_new_resume.type == "application/pdf":
                    resume_text = extract_pdf_text(uploaded_new_resume)
                else:
                    resume_text = extract_docx_text(uploaded_new_resume)
                detected_skills = extract_skills(resume_text)
                updated_skills = ", ".join(detected_skills)
                
            if st.button("Save Changes"):
                if not os.path.exists("profile_pics"):
                    os.makedirs("profile_pics")
                if not os.path.exists("resumes"):
                    os.makedirs("resumes")
                    
                if uploaded_new_image is not None:
                    updated_image_path = os.path.join("profile_pics", f"{employee_id}_{int(time.time())}_{uploaded_new_image.name}")
                    with open(updated_image_path, "wb") as f:
                        f.write(uploaded_new_image.getbuffer())
                        
                if uploaded_new_resume is not None:
                    updated_resume_path = os.path.join("resumes", f"{employee_id}_{int(time.time())}_{uploaded_new_resume.name}")
                    with open(updated_resume_path, "wb") as f:
                        f.write(uploaded_new_resume.getbuffer())
                        
                updated_full_name = f"{new_first_name} {new_last_name}".strip()
                
                cursor.execute('''
                    UPDATE employees 
                    SET name=?, age=?, department=?, email=?, skills=?, image=?, resume=? 
                    WHERE id=?
                ''', (updated_full_name, new_age, new_department, new_email, updated_skills, updated_image_path, updated_resume_path, employee_id))
                conn.commit()
                
                st.success("Information Updated Successfully!")
                time.sleep(1)
                
                st.session_state.open_edit_page = False
                if return_to_role_view in ["admin", "hr"]:
                    st.session_state.page = "Employee Records"
                    st.session_state.updated_successfully = True
                st.rerun()
                
        if st.button("Cancel & Go Back"):
            st.session_state.open_edit_page = False
            st.rerun()

# =================================================
# EMPLOYEE SELF SERVICE DASHBOARD
# =================================================
def employee_self_service():
    st.sidebar.title("Employee Portal")
    st.sidebar.info(f"Connected as: EMPLOYEE")
    if st.sidebar.button("Log Out"):
        logout()
        
    emp_data = fetch_single_employee(st.session_state.employee_id)
    if not emp_data:
        st.error("Error fetching your record profile data.")
        return
          
    st.markdown(f"<h1 class='title'>👋 Welcome, {emp_data[1]}</h1>", unsafe_allow_html=True)
    
    col1, col2 = st.columns([1, 2])
    with col1:
        st.markdown("<div class='box'>", unsafe_allow_html=True)
        st.subheader("Profile Photo")
        if emp_data[6] and os.path.exists(emp_data[6]):
            st.image(emp_data[6], use_container_width=True)
        else:
            st.warning("No picture found.")
        st.markdown("</div>", unsafe_allow_html=True)
        
        if st.button("✏ Edit My Profile Info"):
            st.session_state.edit_employee_id = emp_data[0]
            st.session_state.open_edit_page = True
            st.rerun()
            
    with col2:
        st.markdown("<div class='box'>", unsafe_allow_html=True)
        st.subheader("Personal Credentials & File Registries")
        
        details_df = pd.DataFrame({
            "Metric File Profile": ["Employee ID System", "Full Structural Name", "Age Group", "Allocated Department", "E-Mail Address Connection"],
            "Registered System Value": [emp_data[0], emp_data[1], emp_data[2], emp_data[3], emp_data[4]]
        })
        st.table(details_df.set_index("Metric File Profile"))
        
        st.markdown("### System Logged Core Skills")
        if emp_data[5]:
            skills_list = emp_data[5].split(", ")
            for sk in skills_list:
                st.success(sk)
        else:
            st.info("No system registered keywords recorded yet.")
        st.markdown("</div>", unsafe_allow_html=True)

    if emp_data[7] and os.path.exists(emp_data[7]):
        with st.expander("📄 Click to View My Extracted Resume File Document on System"):
            file_ext = emp_data[7].split(".")[-1]
            if file_ext == "pdf":
                show_pdf(emp_data[7])
            else:
                doc = Document(emp_data[7])
                text = "\n".join([p.text for p in doc.paragraphs])
                st.text_area("Parsed Text Structure", text, height=300)

# =================================================
# ADMINISTRATIVE / HR CONTROL DASHBOARD
# =================================================
def employee_dashboard():
    st.sidebar.title("Navigation")
    if "page" not in st.session_state:
        st.session_state.page = "Dashboard"
        
    menu = st.sidebar.radio(
        "Select Option",
        ["Dashboard", "Employee Registration", "Employee Records"],
        index=["Dashboard", "Employee Registration", "Employee Records"].index(st.session_state.page)
    )
    st.session_state.page = menu
    st.sidebar.success(f"Logged in as {st.session_state.user_role.upper()}")
    
    if st.sidebar.button("Log Out"):
        logout()
        
    if menu == "Dashboard":
        st.markdown("<h1 class='title'>📊 AI HR Dashboard</h1>", unsafe_allow_html=True)
        employees = fetch_employees()
        total_employees = len(employees)
        total_departments = len(set([emp[3] for emp in employees])) if employees else 0
        total_resumes = len([emp for emp in employees if emp[7] != ""])
        
        m1, m2, m3 = st.columns(3)
        m1.metric("Total Employees", total_employees)
        m2.metric("Departments", total_departments)
        m3.metric("Uploaded Resumes", total_resumes)
        st.divider()
        
        if employees:
            df = pd.DataFrame(employees, columns=["ID", "Name", "Age", "Department", "Email", "Skills", "Image", "Resume"])
            st.subheader("Department Wise Employee Count")
            dept_count = df["Department"].value_counts()
            st.bar_chart(dept_count)
            
            st.subheader("Employee Age Distribution")
            age_df = df[["Name", "Age"]].set_index("Name")
            st.line_chart(age_df)
        else:
            st.info("No employee data available")
            
    elif menu == "Employee Registration":
        st.markdown("<h1 class='title'>📝 Employee Registration</h1>", unsafe_allow_html=True)
        col1, col2 = st.columns([1,2])
        
        with col1:
            st.markdown("<div class='box'>", unsafe_allow_html=True)
            st.subheader("Profile Picture")
            uploaded_image = st.file_uploader("Upload JPG/JPEG Image", type=["jpg", "jpeg"])
            image_path = ""
            if uploaded_image is not None:
                image = Image.open(uploaded_image)
                st.image(image, caption="Profile Preview", width=250)
            st.markdown("</div>", unsafe_allow_html=True)
            
        with col2:
            st.markdown("<div class='box'>", unsafe_allow_html=True)
            st.subheader("Employee Details")
            first_name = st.text_input("First Name")
            last_name = st.text_input("Last Name")
            age = st.number_input("Age", min_value=18, max_value=60, step=1)
            department = st.selectbox("Department", ["IT", "HR", "Finance", "Marketing"])
            email = st.text_input("Email")
            
            email_pattern = r'^[\w\.-]+@[\w\.-]+\.\w+$'
            st.subheader("Resume Upload")
            uploaded_resume = st.file_uploader("Upload Resume", type=["pdf", "docx"])
            
            extracted_text = ""
            resume_path = ""
            detected_skills = []
            
            if uploaded_resume is not None:
                if uploaded_resume.type == "application/pdf":
                    extracted_text = extract_pdf_text(uploaded_resume)
                elif uploaded_resume.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    extracted_text = extract_docx_text(uploaded_resume)
                    
                st.text_area("Extracted Resume Text", extracted_text, height=250)
                detected_skills = extract_skills(extracted_text)
                
                st.subheader("Detected Skills")
                if detected_skills:
                    for skill in detected_skills:
                        st.success(skill)
                else:
                    st.warning("No matching skills found")
                    
            if st.button("Register Employee"):
                if first_name == "" or last_name == "" or email == "":
                    st.error("Please fill all required fields")
                elif not re.match(email_pattern, email):
                    st.error("Invalid Email Format")
                else:
                    if not os.path.exists("profile_pics"):
                        os.makedirs("profile_pics")
                    if not os.path.exists("resumes"):
                        os.makedirs("resumes")
                        
                    if uploaded_image is not None:
                        image_path = os.path.join("profile_pics", uploaded_image.name)
                        with open(image_path, "wb") as f:
                            f.write(uploaded_image.getbuffer())
                            
                    if uploaded_resume is not None:
                        resume_path = os.path.join("resumes", uploaded_resume.name)
                        with open(resume_path, "wb") as f:
                            f.write(uploaded_resume.getbuffer())
                            
                    full_name = f"{first_name} {last_name}"
                    cursor.execute('''
                        INSERT INTO employees (name, age, department, email, skills, image, resume)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (full_name, age, department, email, ", ".join(detected_skills), image_path, resume_path))
                    conn.commit()
                    st.success("Employee Registered Successfully")
            st.markdown("</div>", unsafe_allow_html=True)
            
    elif menu == "Employee Records":
        if "updated_successfully" in st.session_state and st.session_state.updated_successfully:
            st.success("Employee details updated successfully")
            st.session_state.updated_successfully = False
            
        st.markdown("<h1 class='title'> Employee Records</h1>", unsafe_allow_html=True)
        employees = fetch_employees()
        
        if not employees:
            st.warning("No employee records found")
        else:
            df = pd.DataFrame(employees, columns=["ID", "Name", "Age", "Department", "Email", "Skills", "Image", "Resume"])
            search = st.text_input(" Search Employee By Name")
            department_filter = st.selectbox("Filter By Department", ["All"] + list(df["Department"].unique()))
            
            filtered_df = df.copy()
            if search:
                filtered_df = filtered_df[filtered_df["Name"].str.contains(search, case=False)]
            if department_filter != "All":
                filtered_df = filtered_df[filtered_df["Department"] == department_filter]
                
            st.subheader("Employee Records Table")
            h1, h2, h3, h4, h5, h6, h7, h8, h9 = st.columns([1,2,1,2,3,2,2,2,2])
            h1.markdown(" ID")
            h2.markdown("Name")
            h3.markdown(" Age")
            h4.markdown(" Dept")
            h5.markdown("Email")
            h6.markdown("Skills")
            h7.markdown(" Profile")
            h8.markdown("Resume")
            h9.markdown(" Action")
            st.divider()
            
            for i, row in filtered_df.iterrows():
                col1, col2, col3, col4, col5, col6, col7, col8, col9 = st.columns([1,2,1,2,3,2,2,2,2])
                col1.write(row["ID"])
                col2.write(row["Name"])
                col3.write(row["Age"])
                col4.write(row["Department"])
                col5.write(row["Email"])
                col6.write(row["Skills"])
                
                with col7:
                    if st.button("View", key=f"profile_{i}"):
                        st.session_state.selected_profile = row["Image"]
                        st.session_state.selected_resume = None
                with col8:
                    if st.button("View", key=f"resume_{i}"):
                        st.session_state.selected_resume = row["Resume"]
                        st.session_state.selected_profile = None
                        
                with col9:
                    edit_col, delete_col = st.columns(2)
                    with edit_col:
                        if st.button("Edit", key=f"edit_{i}"):
                            st.session_state.edit_employee_id = row["ID"]
                            st.session_state.open_edit_page = True
                            st.rerun()
                    with delete_col:
                        if st.button("Del", key=f"delete_{i}"):
                            cursor.execute("DELETE FROM employees WHERE id=?", (row["ID"],))
                            conn.commit()
                            st.success("Deleted")
                            st.rerun()
                st.divider()
                
            if st.session_state.selected_profile:
                st.markdown("##   Profile Picture")
                if os.path.exists(st.session_state.selected_profile):
                    st.image(st.session_state.selected_profile, width=400)
                else:
                    st.error("File not found on system local directory.")
                if st.button("Close Profile View"):
                    st.session_state.selected_profile = None
                    st.rerun()
                    
            if st.session_state.selected_resume:
                st.markdown("## 📄 Resume Viewer")
                if os.path.exists(st.session_state.selected_resume):
                    file_extension = st.session_state.selected_resume.split(".")[-1]
                    if file_extension == "pdf":
                        show_pdf(st.session_state.selected_resume)
                    else:
                        doc = Document(st.session_state.selected_resume)
                        text = "\n".join([p.text for p in doc.paragraphs])
                        st.text_area("Resume Content Text Block", text, height=500)
                else:
                    st.error("File missing from file paths.")  
                if st.button("Close Resume View"):
                    st.session_state.selected_resume = None
                    st.rerun()

# Main Application Runtime Router
if st.session_state.logged_in:
    if st.session_state.open_edit_page:
        render_edit_form(st.session_state.edit_employee_id, return_to_role_view=st.session_state.user_role)
    else:
        if st.session_state.user_role in ["admin", "hr"]:
            employee_dashboard()
        elif st.session_state.user_role == "employee":
            employee_self_service()
else:
    login()

